from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image
from pypdf import PdfReader

from context_gate.report_exports import create_exports, is_export_instruction


def _logo_data_url() -> str:
    payload = io.BytesIO()
    Image.new("RGBA", (60, 40), (255, 0, 255, 255)).save(payload, format="PNG")
    encoded = base64.b64encode(payload.getvalue()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def _state() -> dict[str, object]:
    return {
        "company": {
            "company_name": "Example Company",
            "operator_name": "Demo Operator",
            "important_detail": "Crowd size and event logistics",
            "identity_fields": ["Event name", "Event date"],
            "document_company_header": True,
            "document_footer": "Questions: operations@example.test",
            "company_website": "https://example.test",
            "company_logo_data_url": _logo_data_url(),
            "hidden_sources": [],
            "deleted_sources": [],
        },
        "policy": {"version": "local-v1"},
        "totals": {"total": 3, "allow": 1, "review": 1, "block": 1},
        "cases": [
            {
                "case_id": "A1",
                "title": "Verified organizer update",
                "original_outcome": "ALLOW",
                "outcome": "ALLOW",
                "source": "Official email",
                "summary": "Current verified evidence supports the proposed value.",
                "evidence_count": 1,
                "rule_ids": ["CG-001"],
                "corrected": False,
            },
            {
                "case_id": "R1",
                "title": "Missing provenance",
                "original_outcome": "REVIEW",
                "outcome": "REVIEW",
                "source": "Unknown",
                "summary": "A person must confirm the source before use.",
                "evidence_count": 1,
                "rule_ids": ["CG-003"],
                "corrected": False,
            },
            {
                "case_id": "B1",
                "title": "Lower-authority conflict",
                "original_outcome": "BLOCK",
                "outcome": "BLOCK",
                "source": "Community listing, official email",
                "summary": "The requested value conflicts with stronger evidence.",
                "evidence_count": 2,
                "rule_ids": ["CG-002"],
                "corrected": False,
            },
        ],
        "patterns": [
            {
                "label": "Crowd-size update",
                "count": 113,
                "description": "35 confirmed plus 78 more equals 113.",
            }
        ],
        "source_summary": {
            "fictional": True,
            "messages_scanned": 10,
            "distinct_events": 9,
        },
        "connectors": {"secret_token": "must-not-appear"},
    }


def test_combined_chat_export_creates_valid_docx_pdf_and_png(tmp_path: Path) -> None:
    artifacts = create_exports(
        _state(),
        "Create a report and a pie chart that I can send out.",
        output_root=tmp_path,
    )
    by_kind = {item.kind: Path(item.path) for item in artifacts}

    assert set(by_kind) == {"docx", "pdf", "png"}
    assert all(path.parent == tmp_path for path in by_kind.values())
    assert all(path.stat().st_size > 1_000 for path in by_kind.values())

    word = Document(by_kind["docx"])
    assert word.paragraphs[0].text == "Example Company"
    word_text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    word_text += "\n" + "\n".join(
        cell.text for table in word.tables for row in table.rows for cell in row.cells
    )
    assert "ContextGate Evidence Report" in word_text
    assert "Example Company" in word_text
    assert "https://example.test" in word_text
    assert "operations@example.test" in "\n".join(
        paragraph.text
        for section in word.sections
        for paragraph in section.footer.paragraphs
    )
    assert "must-not-appear" not in word_text
    with zipfile.ZipFile(by_kind["docx"]) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())

    pdf = PdfReader(str(by_kind["pdf"]))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert len(pdf.pages) >= 2
    assert "ContextGate Evidence Report" in pdf_text
    assert "Example Company" in pdf_text
    assert "https://example.test" in pdf_text
    assert "operations@example.test" in pdf_text
    assert "must-not-appear" not in pdf_text

    with Image.open(by_kind["png"]) as chart:
        assert chart.size == (1280, 720)
        assert chart.format == "PNG"
        assert chart.getpixel((1_130, 135)) == (255, 0, 255)


def test_explicit_export_formats_and_instruction_detection(tmp_path: Path) -> None:
    assert is_export_instruction("Make a graph of what you are showing me")
    assert is_export_instruction("Save a Word report")
    assert is_export_instruction("Please prepare a report and graph")
    assert not is_export_instruction("What does this report mean?")
    assert not is_export_instruction("Don't create a report")
    assert not is_export_instruction("Do not save a PDF")

    chart = create_exports(_state(), "Create a bar graph", output_root=tmp_path)
    assert [item.kind for item in chart] == ["png"]

    report = create_exports(_state(), "Make an HTML report", output_root=tmp_path)
    assert [item.kind for item in report] == ["html"]
    html_report = Path(report[0].path).read_text(encoding="utf-8")
    assert "ContextGate Evidence Report" in html_report
    assert "https://example.test" in html_report
    assert "data:image/png;base64," in html_report

    anonymous = create_exports(
        _state(),
        "Make an HTML report without the company name",
        output_root=tmp_path,
    )
    anonymous_text = Path(anonymous[0].path).read_text(encoding="utf-8")
    assert "Example Company" not in anonymous_text
    assert "operations@example.test" in anonymous_text


def test_exports_allowlist_excludes_secrets_from_every_output(tmp_path: Path) -> None:
    state = _state()
    marker = "CG-NEVER-EXPORT-client-secret-access-token"
    state["connectors"] = {
        "google": {
            "client_secret": marker,
            "accounts": [{"access_token": marker}],
        }
    }
    state["chat_history"] = [{"role": "user", "text": marker}]
    state["source_status"] = {"provider_debug_details": marker}
    state["company"]["profile_error"] = marker

    artifacts = create_exports(
        state,
        "Create a PDF Word HTML report and PNG chart",
        output_root=tmp_path,
    )
    by_kind = {item.kind: Path(item.path) for item in artifacts}
    assert set(by_kind) == {"docx", "pdf", "html", "png"}

    with zipfile.ZipFile(by_kind["docx"]) as archive:
        docx_xml = b"\n".join(
            archive.read(name)
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        ).decode("utf-8", errors="replace")
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(str(by_kind["pdf"])).pages
    )
    html_text = by_kind["html"].read_text(encoding="utf-8")
    with Image.open(by_kind["png"]) as image:
        png_metadata = " ".join(str(value) for value in image.info.values())

    assert marker not in docx_xml
    assert marker not in pdf_text
    assert marker not in html_text
    assert marker not in png_metadata
    assert all(marker not in artifact.filename for artifact in artifacts)


def test_company_name_suppression_applies_to_docx_pdf_html_and_filename(
    tmp_path: Path,
) -> None:
    artifacts = create_exports(
        _state(),
        "Create a PDF Word HTML report without company name",
        output_root=tmp_path,
    )
    by_kind = {item.kind: Path(item.path) for item in artifacts}

    assert set(by_kind) == {"docx", "pdf", "html"}
    assert all("contextgate-workspace-" in item.filename for item in artifacts)
    word = Document(by_kind["docx"])
    word_text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    word_text += "\n" + "\n".join(
        paragraph.text
        for section in word.sections
        for paragraph in section.header.paragraphs
    )
    word_text += "\n" + "\n".join(
        paragraph.text
        for section in word.sections
        for paragraph in section.footer.paragraphs
    )
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(str(by_kind["pdf"])).pages
    )
    html_text = by_kind["html"].read_text(encoding="utf-8")

    for rendered in (word_text, pdf_text, html_text):
        assert "Example Company" not in rendered
        assert "https://example.test" not in rendered
        assert "ContextGate Evidence Report" in rendered
        assert "operations@example.test" in rendered


def test_saved_header_preference_suppresses_company_name_by_default(
    tmp_path: Path,
) -> None:
    state = _state()
    state["company"]["document_company_header"] = False
    artifact = create_exports(state, "Make an HTML report", output_root=tmp_path)[0]
    rendered = Path(artifact.path).read_text(encoding="utf-8")

    assert "Example Company" not in rendered
    assert "https://example.test" not in rendered
    assert "contextgate-workspace-" in artifact.filename


def test_grouped_metric_exports_include_totals_and_row_evidence(tmp_path: Path) -> None:
    state = _state()
    state["tracking"] = {
        "active_topic_id": "topic-sales",
        "topics": [{"topic_id": "topic-sales", "name": "office sales"}],
    }
    state["grouped_metrics"] = [
        {
            "dataset_id": "metric-sales",
            "dataset_name": "fictional office sales",
            "topic_id": "topic-sales",
            "topic_name": "office sales",
            "source_filename": "fictional-office-sales.csv",
            "source_reference": "upload://sales-evidence",
            "group_field": "office",
            "metric_field": "sales",
            "unit": None,
            "row_count": 4,
            "group_totals": {"New York": 189, "Austin": 73},
            "evidence": [
                {
                    "row_number": 2,
                    "group": "New York",
                    "value": 100,
                    "reference": "upload://sales-evidence#row=2",
                },
                {
                    "row_number": 3,
                    "group": "Austin",
                    "value": 40,
                    "reference": "upload://sales-evidence#row=3",
                },
                {
                    "row_number": 4,
                    "group": "New York",
                    "value": 89,
                    "reference": "upload://sales-evidence#row=4",
                },
                {
                    "row_number": 5,
                    "group": "Austin",
                    "value": 33,
                    "reference": "upload://sales-evidence#row=5",
                },
            ],
            "fictional": True,
        }
    ]
    artifacts = create_exports(
        state,
        "Create a PDF Word HTML report and pie chart",
        output_root=tmp_path,
    )
    by_kind = {item.kind: Path(item.path) for item in artifacts}

    word = Document(by_kind["docx"])
    word_text = "\n".join(
        [*(paragraph.text for paragraph in word.paragraphs)]
        + [
            cell.text
            for table in word.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(str(by_kind["pdf"])).pages
    )
    html_text = by_kind["html"].read_text(encoding="utf-8")
    for rendered in (word_text, pdf_text, html_text):
        assert "fictional office sales" in rendered
        assert "New York" in rendered
        assert "189" in rendered
        assert "upload://sales-evidence#row=2" in rendered
    with Image.open(by_kind["png"]) as chart:
        assert chart.size == (1280, 720)
