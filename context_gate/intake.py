"""Bounded, local-only intake for untrusted evidence artifacts.

The intake boundary deliberately separates *extracting text* from *trusting a
claim*.  Receipts never contain the original bytes, and candidates created
from a receipt remain unverified until another policy step establishes trust.
"""

from __future__ import annotations

import hashlib
import io
import re
import unicodedata
from datetime import datetime
from email import policy
from email.parser import BytesParser
from enum import StrEnum
from html.parser import HTMLParser
from pathlib import PurePosixPath
from typing import ClassVar, Final

from pydantic import Field, model_validator

from .models import (
    ContextEvent,
    EvidenceStatus,
    Sensitivity,
    StrictModel,
)
from .normalization import normalize_event

MAX_ARTIFACT_BYTES: Final = 10 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS: Final = 20_000
MAX_SAFE_FILENAME_CHARS: Final = 180


class ArtifactStatus(StrEnum):
    EXTRACTED = "EXTRACTED"
    NO_TEXT = "NO_TEXT"
    OCR_REQUIRED = "OCR_REQUIRED"
    DEPENDENCY_REQUIRED = "DEPENDENCY_REQUIRED"
    UNSUPPORTED = "UNSUPPORTED"
    EXTRACTION_FAILED = "EXTRACTION_FAILED"


class ArtifactIntakeError(ValueError):
    """A safe, caller-visible rejection with no artifact content attached."""


class ArtifactReceipt(StrictModel):
    artifact_id: str = Field(pattern=r"^artifact-[a-f0-9]{24}$")
    safe_filename: str = Field(min_length=1, max_length=MAX_SAFE_FILENAME_CHARS)
    content_type: str = Field(min_length=1, max_length=255)
    size_bytes: int = Field(ge=0, le=MAX_ARTIFACT_BYTES)
    sha256: str = Field(pattern=r"^[a-f0-9]{64}$")
    status: ArtifactStatus
    extracted_text: str | None = Field(
        default=None, max_length=MAX_EXTRACTED_TEXT_CHARS
    )
    extracted_chars: int = Field(ge=0, le=MAX_EXTRACTED_TEXT_CHARS)
    truncated: bool = False
    extractor: str = Field(min_length=1, max_length=64)
    message: str = Field(min_length=1, max_length=512)

    @model_validator(mode="after")
    def extracted_length_is_consistent(self) -> ArtifactReceipt:
        actual = len(self.extracted_text or "")
        if self.extracted_chars != actual:
            raise ValueError("extracted_chars must match extracted_text")
        if self.status == ArtifactStatus.EXTRACTED and not self.extracted_text:
            raise ValueError("EXTRACTED receipts must contain text")
        if self.status != ArtifactStatus.EXTRACTED and self.extracted_text is not None:
            raise ValueError("only EXTRACTED receipts may contain text")
        return self


class _VisibleTextParser(HTMLParser):
    """Small, non-networking HTML/XML text extractor."""

    _ignored: ClassVar[set[str]] = {"script", "style", "template", "noscript"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._ignore_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        if tag.casefold() in self._ignored:
            self._ignore_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in self._ignored and self._ignore_depth:
            self._ignore_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignore_depth:
            self.parts.append(data)


_TEXT_MIME_TYPES: Final = {
    "application/csv",
    "application/json",
    "application/markdown",
    "application/xhtml+xml",
    "application/xml",
    "text/csv",
    "text/html",
    "text/json",
    "text/markdown",
    "text/plain",
    "text/xml",
}
_TEXT_SUFFIXES: Final = {
    ".csv",
    ".htm",
    ".html",
    ".json",
    ".markdown",
    ".md",
    ".txt",
    ".xml",
}
_MARKUP_MIME_TYPES: Final = {
    "application/xhtml+xml",
    "application/xml",
    "text/html",
    "text/xml",
}
_MARKUP_SUFFIXES: Final = {".htm", ".html", ".xml"}
_IMAGE_MIME_TYPES: Final = {
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/webp",
}
_IMAGE_SUFFIXES: Final = {".gif", ".jpeg", ".jpg", ".png", ".webp"}
_EMAIL_MIME_TYPES: Final = {"message/rfc822"}
_EMAIL_SUFFIXES: Final = {".eml"}


def safe_basename(filename: str) -> str:
    """Return a bounded, portable basename without path components."""

    if not isinstance(filename, str):
        raise ArtifactIntakeError("filename must be text")
    if len(filename) > 4096:
        raise ArtifactIntakeError("filename exceeds the 4096-character limit")
    normalized = unicodedata.normalize("NFKC", filename).replace("\\", "/")
    basename = PurePosixPath(normalized).name
    basename = re.sub(r"[^A-Za-z0-9._-]+", "_", basename).strip(" ._")
    if not basename or basename in {".", ".."}:
        basename = "upload.bin"

    if len(basename) > MAX_SAFE_FILENAME_CHARS:
        suffix = PurePosixPath(basename).suffix[:20]
        stem_limit = MAX_SAFE_FILENAME_CHARS - len(suffix)
        basename = f"{basename[:stem_limit]}{suffix}"
    return basename


def _normal_content_type(content_type: str | None) -> str:
    if content_type is None:
        return "application/octet-stream"
    if not isinstance(content_type, str):
        raise ArtifactIntakeError("content_type must be text")
    if len(content_type) > 255:
        raise ArtifactIntakeError("content_type exceeds the 255-character limit")
    if not content_type.strip():
        return "application/octet-stream"
    normalized = content_type.split(";", 1)[0].strip().casefold()
    if len(normalized) > 255 or not re.fullmatch(
        r"[a-z0-9!#$&^_.+-]+/[a-z0-9!#$&^_.+-]+", normalized
    ):
        raise ArtifactIntakeError("content_type must be a valid bounded media type")
    return normalized


def _bounded_text(text: str) -> tuple[str, bool]:
    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if len(normalized) <= MAX_EXTRACTED_TEXT_CHARS:
        return normalized, False
    return normalized[:MAX_EXTRACTED_TEXT_CHARS], True


def _is_probably_binary(content: bytes) -> bool:
    sample = content[:4096]
    if b"\x00" in sample:
        return True
    if not sample:
        return False
    disallowed = sum(byte < 32 and byte not in {9, 10, 12, 13} for byte in sample)
    return disallowed / len(sample) > 0.05


def _text_receipt_fields(
    content: bytes, *, strip_markup: bool
) -> tuple[ArtifactStatus, str | None, bool, str]:
    if _is_probably_binary(content):
        return ArtifactStatus.UNSUPPORTED, None, False, "Content appears to be binary."
    try:
        decoded = content.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError:
        return (
            ArtifactStatus.UNSUPPORTED,
            None,
            False,
            "Content is not valid UTF-8 text.",
        )

    if strip_markup:
        parser = _VisibleTextParser()
        try:
            parser.feed(decoded)
            parser.close()
        except Exception:  # noqa: BLE001 - untrusted parsers must fail closed
            return (
                ArtifactStatus.EXTRACTION_FAILED,
                None,
                False,
                "Markup text extraction failed safely.",
            )
        decoded = " ".join(parser.parts)

    text, truncated = _bounded_text(decoded)
    if not text:
        return ArtifactStatus.NO_TEXT, None, False, "No extractable text was found."
    return ArtifactStatus.EXTRACTED, text, truncated, "Text extracted locally."


def _pdf_receipt_fields(
    content: bytes,
) -> tuple[ArtifactStatus, str | None, bool, str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return (
            ArtifactStatus.DEPENDENCY_REQUIRED,
            None,
            False,
            "Install pypdf to extract text-layer PDFs.",
            "pypdf-unavailable",
        )

    try:
        reader = PdfReader(io.BytesIO(content))
        parts = [(page.extract_text() or "") for page in reader.pages]
    except Exception:  # noqa: BLE001 - optional parser errors vary by version
        return (
            ArtifactStatus.EXTRACTION_FAILED,
            None,
            False,
            "PDF text extraction failed safely.",
            "pypdf",
        )
    text, truncated = _bounded_text("\n".join(parts))
    if not text:
        return (
            ArtifactStatus.OCR_REQUIRED,
            None,
            False,
            "The PDF has no text layer; OCR is required.",
            "pypdf",
        )
    return (
        ArtifactStatus.EXTRACTED,
        text,
        truncated,
        "PDF text layer extracted locally.",
        "pypdf",
    )


def _docx_receipt_fields(
    content: bytes,
) -> tuple[ArtifactStatus, str | None, bool, str, str]:
    try:
        from docx import Document
    except ImportError:
        return (
            ArtifactStatus.DEPENDENCY_REQUIRED,
            None,
            False,
            "Install python-docx to extract DOCX text.",
            "python-docx-unavailable",
        )

    try:
        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts.extend(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
    except Exception:  # noqa: BLE001 - optional parser errors vary by version
        return (
            ArtifactStatus.EXTRACTION_FAILED,
            None,
            False,
            "DOCX text extraction failed safely.",
            "python-docx",
        )
    text, truncated = _bounded_text("\n".join(parts))
    if not text:
        return (
            ArtifactStatus.NO_TEXT,
            None,
            False,
            "No extractable DOCX text was found.",
            "python-docx",
        )
    return (
        ArtifactStatus.EXTRACTED,
        text,
        truncated,
        "DOCX text extracted locally.",
        "python-docx",
    )


def _email_receipt_fields(
    content: bytes,
) -> tuple[ArtifactStatus, str | None, bool, str, str]:
    """Extract subject and visible text parts without retaining MIME attachments."""

    try:
        message = BytesParser(policy=policy.default).parsebytes(content)
        parts = []
        subject = str(message.get("subject") or "").strip()
        if subject:
            parts.append(f"Subject: {subject}")
        for part in message.walk():
            if part.is_multipart() or part.get_content_disposition() == "attachment":
                continue
            media_type = part.get_content_type().casefold()
            if media_type not in {"text/plain", "text/html"}:
                continue
            body = part.get_content()
            if not isinstance(body, str):
                continue
            if media_type == "text/html":
                parser = _VisibleTextParser()
                parser.feed(body)
                parser.close()
                body = " ".join(parser.parts)
            if body.strip():
                parts.append(body)
    except Exception:  # noqa: BLE001 - malformed MIME must fail closed
        return (
            ArtifactStatus.EXTRACTION_FAILED,
            None,
            False,
            "Email text extraction failed safely.",
            "stdlib-email",
        )
    text, truncated = _bounded_text("\n".join(parts))
    if not text:
        return (
            ArtifactStatus.NO_TEXT,
            None,
            False,
            "No extractable email text was found.",
            "stdlib-email",
        )
    return (
        ArtifactStatus.EXTRACTED,
        text,
        truncated,
        "Email subject and visible text extracted locally; attachments were skipped.",
        "stdlib-email",
    )


def ingest_artifact(
    filename: str,
    content_type: str | None,
    content: bytes | bytearray | memoryview,
    *,
    max_bytes: int = MAX_ARTIFACT_BYTES,
) -> ArtifactReceipt:
    """Hash and inspect an artifact in memory without persisting its bytes."""

    if not isinstance(max_bytes, int) or isinstance(max_bytes, bool):
        raise ArtifactIntakeError("max_bytes must be an integer")
    if max_bytes < 1 or max_bytes > MAX_ARTIFACT_BYTES:
        raise ArtifactIntakeError(
            f"max_bytes must be between 1 and {MAX_ARTIFACT_BYTES}"
        )
    if not isinstance(content, (bytes, bytearray, memoryview)):
        raise ArtifactIntakeError("content must be bytes")
    payload = bytes(content)
    if len(payload) > max_bytes:
        raise ArtifactIntakeError(f"artifact exceeds the {max_bytes}-byte limit")

    safe_filename = safe_basename(filename)
    media_type = _normal_content_type(content_type)
    suffix = PurePosixPath(safe_filename).suffix.casefold()
    digest = hashlib.sha256(payload).hexdigest()

    looks_pdf = (
        media_type == "application/pdf"
        or suffix == ".pdf"
        or payload.startswith(b"%PDF-")
    )
    looks_docx = (
        media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == ".docx"
    )
    has_image_signature = payload.startswith(
        (b"\x89PNG\r\n\x1a\n", b"\xff\xd8\xff", b"GIF87a", b"GIF89a")
    ) or (payload.startswith(b"RIFF") and payload[8:12] == b"WEBP")
    looks_image = (
        media_type in _IMAGE_MIME_TYPES
        or suffix in _IMAGE_SUFFIXES
        or has_image_signature
    )
    looks_email = media_type in _EMAIL_MIME_TYPES or suffix in _EMAIL_SUFFIXES
    looks_text = media_type in _TEXT_MIME_TYPES or suffix in _TEXT_SUFFIXES

    extractor = "none"
    if looks_pdf:
        status, text, truncated, message, extractor = _pdf_receipt_fields(payload)
    elif looks_docx:
        status, text, truncated, message, extractor = _docx_receipt_fields(payload)
    elif looks_email:
        status, text, truncated, message, extractor = _email_receipt_fields(payload)
    elif looks_image:
        status, text, truncated, message, extractor = (
            ArtifactStatus.OCR_REQUIRED,
            None,
            False,
            "Image text requires an explicit OCR step; no text was inferred.",
            "ocr-not-run",
        )
    elif looks_text:
        status, text, truncated, message = _text_receipt_fields(
            payload,
            strip_markup=(
                media_type in _MARKUP_MIME_TYPES or suffix in _MARKUP_SUFFIXES
            ),
        )
        extractor = (
            "utf-8-markup"
            if (media_type in _MARKUP_MIME_TYPES or suffix in _MARKUP_SUFFIXES)
            else "utf-8"
        )
    else:
        status, text, truncated, message, extractor = (
            ArtifactStatus.UNSUPPORTED,
            None,
            False,
            "This artifact type is not supported for local text extraction.",
            "none",
        )

    return ArtifactReceipt(
        artifact_id=f"artifact-{digest[:24]}",
        safe_filename=safe_filename,
        content_type=media_type,
        size_bytes=len(payload),
        sha256=digest,
        status=status,
        extracted_text=text,
        extracted_chars=len(text or ""),
        truncated=truncated,
        extractor=extractor,
        message=message,
    )


def create_context_event_candidate(
    receipt: ArtifactReceipt,
    *,
    event_id: str,
    entity_id: str,
    field_name: str,
    claim_value: str,
    source_name: str,
    source_type: str,
    trust_score: float = 0.0,
    observed_at: datetime | None = None,
    effective_at: datetime | None = None,
    sensitivity: Sensitivity = Sensitivity.PUBLIC,
) -> ContextEvent:
    """Create an explicitly unverified claim linked to an immutable upload digest.

    ``claim_value`` and source fields come from the caller, not extracted text.
    The extraction receipt therefore proves only which bytes were inspected; it
    cannot authenticate either the claim or its source.
    """

    candidate = ContextEvent(
        event_id=event_id,
        entity_id=entity_id,
        field_name=field_name,
        field_value=claim_value,
        source_name=source_name,
        source_type=source_type,
        trust_score=trust_score,
        observed_at=observed_at,
        effective_at=effective_at,
        sensitivity=sensitivity,
        evidence_reference=f"upload-sha256://{receipt.sha256}",
        status=EvidenceStatus.UNVERIFIED,
    )
    return normalize_event(candidate)
