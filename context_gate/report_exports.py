"""Local, evidence-labelled report and chart exports for the command center."""

from __future__ import annotations

import base64
import binascii
import html
import io
import json
import math
import os
import re
import tempfile
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal
from uuid import uuid4

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

MAX_EXPORT_CASES = 100
MAX_EXPORT_PATTERNS = 100
MAX_EXPORT_METRIC_DATASETS = 12
MAX_EXPORT_METRIC_GROUPS = 50
MAX_EXPORT_METRIC_EVIDENCE = 500


class ReportExportError(ValueError):
    """Raised with a safe local-export error message."""


@dataclass(frozen=True)
class ExportArtifact:
    """One generated file returned to the chat boundary."""

    kind: str
    path: str
    filename: str
    size_bytes: int

    def model_dump(self) -> dict[str, object]:
        return asdict(self)


def _text(value: object, fallback: str = "Not supplied", limit: int = 2_000) -> str:
    if value is None or value == "":
        return fallback
    if isinstance(value, (str, int, float, bool)):
        cleaned = " ".join(str(value).split())
    else:
        cleaned = " ".join(json.dumps(value, ensure_ascii=False).split())
    return cleaned[:limit]


def _grouped_metric_snapshot(state: dict[str, Any]) -> list[dict[str, Any]]:
    datasets = state.get("grouped_metrics")
    if not isinstance(datasets, list):
        return []
    normalized: list[dict[str, Any]] = []
    for item in datasets[:MAX_EXPORT_METRIC_DATASETS]:
        if not isinstance(item, dict) or not isinstance(item.get("group_totals"), dict):
            continue
        totals: dict[str, int | float] = {}
        for group, value in list(item["group_totals"].items())[
            :MAX_EXPORT_METRIC_GROUPS
        ]:
            if (
                isinstance(value, bool)
                or not isinstance(value, (int, float))
                or not math.isfinite(value)
            ):
                continue
            totals[_text(group, "Unnamed group", 200)] = value
        evidence = []
        raw_evidence = item.get("evidence")
        if isinstance(raw_evidence, list):
            for row in raw_evidence[:MAX_EXPORT_METRIC_EVIDENCE]:
                if not isinstance(row, dict):
                    continue
                evidence.append(
                    {
                        "row_number": int(row.get("row_number") or 0),
                        "group": _text(row.get("group"), "Unnamed group", 200),
                        "value": row.get("value"),
                        "reference": _text(row.get("reference"), "Not supplied", 1_000),
                    }
                )
        if totals:
            normalized.append(
                {
                    "dataset_id": _text(item.get("dataset_id"), "metric", 100),
                    "dataset_name": _text(
                        item.get("dataset_name"), "Grouped metric", 160
                    ),
                    "topic_id": _text(item.get("topic_id"), "", 100),
                    "topic_name": _text(item.get("topic_name"), "", 120),
                    "source_filename": _text(
                        item.get("source_filename"), "Uploaded data", 512
                    ),
                    "source_reference": _text(
                        item.get("source_reference"), "Not supplied", 1_000
                    ),
                    "group_field": _text(item.get("group_field"), "group", 100),
                    "metric_field": _text(item.get("metric_field"), "metric", 100),
                    "unit": _text(item.get("unit"), "", 40),
                    "row_count": int(item.get("row_count") or len(evidence)),
                    "group_totals": totals,
                    "evidence": evidence,
                    "fictional": bool(item.get("fictional")),
                }
            )
    return normalized


def _normalized_snapshot(state: dict[str, Any]) -> dict[str, Any]:
    company = state.get("company") if isinstance(state.get("company"), dict) else {}
    totals = state.get("totals") if isinstance(state.get("totals"), dict) else {}
    cases = state.get("cases") if isinstance(state.get("cases"), list) else []
    patterns = state.get("patterns") if isinstance(state.get("patterns"), list) else []
    source_summary = (
        state.get("source_summary")
        if isinstance(state.get("source_summary"), dict)
        else {}
    )
    safe_cases = []
    for item in cases[:MAX_EXPORT_CASES]:
        if not isinstance(item, dict):
            continue
        safe_cases.append(
            {
                "case_id": _text(item.get("case_id"), "—", 24),
                "title": _text(item.get("title") or item.get("name"), "Case", 160),
                "original_outcome": _text(
                    item.get("original_outcome") or item.get("outcome"), "UNKNOWN", 16
                ).upper(),
                "effective_outcome": _text(item.get("outcome"), "UNKNOWN", 16).upper(),
                "source": _text(item.get("source"), "Local evidence", 240),
                "summary": _text(
                    item.get("summary") or item.get("description"),
                    "No summary supplied.",
                    600,
                ),
                "evidence_count": int(item.get("evidence_count") or 0),
                "rule_ids": [
                    _text(rule, "", 120) for rule in (item.get("rule_ids") or [])[:12]
                ],
                "corrected": bool(item.get("corrected")),
            }
        )
    grouped_metrics = _grouped_metric_snapshot(state)
    tracking = state.get("tracking") if isinstance(state.get("tracking"), dict) else {}
    return {
        "generated_at": datetime.now(UTC).isoformat(),
        "company_name": _text(company.get("company_name"), "Company workspace", 80),
        "operator_name": _text(company.get("operator_name"), "Local operator", 80),
        "show_company_header": bool(company.get("document_company_header", True)),
        "document_footer": _text(company.get("document_footer"), "", 240),
        "company_website": _text(company.get("company_website"), "", 240),
        "company_logo_data_url": (
            company.get("company_logo_data_url")
            if isinstance(company.get("company_logo_data_url"), str)
            and str(company.get("company_logo_data_url")).startswith(
                "data:image/png;base64,"
            )
            and len(str(company.get("company_logo_data_url"))) <= 1_500_000
            else ""
        ),
        "important_detail": _text(
            company.get("important_detail"), "Not configured", 120
        ),
        "identity_fields": [
            _text(item, "", 80) for item in (company.get("identity_fields") or [])[:8]
        ],
        "hidden_sources": [
            _text(item, "", 120) for item in (company.get("hidden_sources") or [])[:32]
        ],
        "deleted_sources": [
            _text(item, "", 120) for item in (company.get("deleted_sources") or [])[:32]
        ],
        "policy_version": _text(
            (state.get("policy") or {}).get("version")
            if isinstance(state.get("policy"), dict)
            else None,
            "active",
            80,
        ),
        "totals": {
            "TOTAL": int(totals.get("total") or len(safe_cases)),
            "ALLOW": int(totals.get("allow") or 0),
            "REVIEW": int(totals.get("review") or 0),
            "BLOCK": int(totals.get("block") or 0),
        },
        "cases": safe_cases,
        "patterns": [
            item for item in patterns[:MAX_EXPORT_PATTERNS] if isinstance(item, dict)
        ],
        "grouped_metrics": grouped_metrics,
        "tracking": {
            "active_topic_id": _text(tracking.get("active_topic_id"), "", 100),
            "topics": [
                item
                for item in (tracking.get("topics") or [])[:24]
                if isinstance(item, dict)
            ],
        },
        "source_summary": source_summary,
        "fictional": bool(source_summary.get("fictional", False)),
        "action_boundary": "No external action or email was executed by this export.",
    }


def _output_root(output_root: str | os.PathLike[str] | None) -> Path:
    if output_root is not None:
        root = Path(output_root)
    else:
        documents = Path.home() / "Documents"
        root = documents / "ContextGate Exports"
    try:
        root.mkdir(parents=True, exist_ok=True)
        if not root.is_dir():
            raise OSError
        return root.resolve()
    except OSError:
        raise ReportExportError(
            "The Documents export folder is not writable."
        ) from None


def _atomic_write(path: Path, content: bytes) -> None:
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{path.stem}-", suffix=".tmp", dir=path.parent
    )
    try:
        with os.fdopen(descriptor, "wb") as stream:
            stream.write(content)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary_name, path)
    except Exception:
        try:
            os.unlink(temporary_name)
        except OSError:
            pass
        raise


def _font(size: int, *, bold: bool = False, mono: bool = False) -> ImageFont.ImageFont:
    candidates = (
        [Path("C:/Windows/Fonts/consolab.ttf"), Path("C:/Windows/Fonts/consola.ttf")]
        if mono
        else [
            Path("C:/Windows/Fonts/seguisb.ttf")
            if bold
            else Path("C:/Windows/Fonts/segoeui.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
            if bold
            else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]
    )
    for path in candidates:
        if path.is_file():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _logo_bytes(snapshot: dict[str, Any]) -> bytes | None:
    data_url = snapshot.get("company_logo_data_url")
    if not isinstance(data_url, str) or not data_url.startswith(
        "data:image/png;base64,"
    ):
        return None
    try:
        content = base64.b64decode(data_url.split(",", 1)[1], validate=True)
    except (binascii.Error, ValueError):
        return None
    return content if 0 < len(content) <= 1_000_000 else None


def _active_metric_dataset(snapshot: dict[str, Any]) -> dict[str, Any] | None:
    datasets = snapshot.get("grouped_metrics") or []
    active_topic_id = (snapshot.get("tracking") or {}).get("active_topic_id")
    if active_topic_id:
        match = next(
            (
                item
                for item in reversed(datasets)
                if item.get("topic_id") == active_topic_id
            ),
            None,
        )
        if match is not None:
            return match
    return datasets[-1] if datasets else None


def _chart_png(snapshot: dict[str, Any], chart_type: Literal["pie", "bar"]) -> bytes:
    image = Image.new("RGB", (1280, 720), "#03070d")
    draw = ImageDraw.Draw(image)
    for x in range(0, 1280, 40):
        draw.line((x, 0, x, 720), fill="#0c2638", width=1)
    for y in range(0, 720, 40):
        draw.line((0, y, 1280, y), fill="#0c2638", width=1)
    draw.rounded_rectangle(
        (42, 38, 1238, 680), 22, fill="#07131f", outline="#23789d", width=2
    )
    if snapshot["show_company_header"]:
        draw.text(
            (82, 70),
            snapshot["company_name"],
            font=_font(35, bold=True),
            fill="#f2fbff",
        )
        product_y = 122
    else:
        product_y = 82
    draw.text(
        (82, product_y),
        "CONTEXTGATE / DECISION INTELLIGENCE",
        font=_font(21, mono=True),
        fill="#38d8ff",
    )
    metadata_y = 163
    if snapshot["company_website"]:
        draw.text(
            (82, 158),
            snapshot["company_website"],
            font=_font(16),
            fill="#82e8ff",
        )
        metadata_y = 187
    draw.text(
        (82, metadata_y),
        f"Operator: {snapshot['operator_name']}  |  Generated {snapshot['generated_at'][:19]} UTC",
        font=_font(17),
        fill="#8ba2b4",
    )
    metric_dataset = _active_metric_dataset(snapshot)
    fictional = (
        bool(metric_dataset.get("fictional"))
        if metric_dataset is not None
        else snapshot["fictional"]
    )
    origin = "FICTIONAL DEMO" if fictional else "SCANNED / UPLOADED SOURCES"
    draw.rounded_rectangle(
        (962, 76, 1196, 116), 9, fill="#0b2c31", outline="#39e6a0", width=1
    )
    draw.text((984, 87), origin, font=_font(14, bold=True, mono=True), fill="#39e6a0")
    logo_content = _logo_bytes(snapshot)
    if logo_content:
        try:
            with Image.open(io.BytesIO(logo_content)) as logo_source:
                logo = logo_source.convert("RGBA")
                logo.thumbnail((130, 82), Image.Resampling.LANCZOS)
                logo_x = 1188 - logo.width
                image.paste(logo, (logo_x, 132), logo)
        except (OSError, ValueError):
            pass

    if metric_dataset is not None:
        metric = str(metric_dataset["metric_field"])
        group_field = str(metric_dataset["group_field"])
        rows = list(metric_dataset["group_totals"].items())[:6]
        palette = ["#38d8ff", "#39e6a0", "#ffc857", "#b985ff", "#ff8f70", "#78a8ff"]
        draw.text(
            (82, 226),
            f"{metric.upper()} BY {group_field.upper()}",
            font=_font(20, bold=True, mono=True),
            fill="#d5e5ef",
        )
        draw.text(
            (82, 258),
            str(metric_dataset["dataset_name"]),
            font=_font(16),
            fill="#8ba2b4",
        )
        values = [float(value) for _, value in rows]
        if (
            chart_type == "pie"
            and rows
            and all(value >= 0 for value in values)
            and sum(values) > 0
        ):
            box = (90, 292, 450, 652)
            start = -90.0
            denominator = sum(values)
            for index, value in enumerate(values):
                end = start + 360 * value / denominator
                draw.pieslice(box, start=start, end=end, fill=palette[index])
                start = end
            draw.ellipse((190, 392, 350, 552), fill="#07131f")
            for index, (group, value) in enumerate(rows):
                y = 304 + index * 52
                draw.ellipse((540, y + 8, 560, y + 28), fill=palette[index])
                draw.text((578, y), str(group)[:34], font=_font(18), fill="#93acbd")
                unit = f" {metric_dataset['unit']}" if metric_dataset["unit"] else ""
                draw.text(
                    (1035, y),
                    f"{value}{unit}",
                    font=_font(21, bold=True),
                    fill="#f2fbff",
                )
        else:
            maximum = max((abs(value) for value in values), default=1.0) or 1.0
            for index, (group, value) in enumerate(rows):
                y = 302 + index * 54
                numeric = float(value)
                draw.text((82, y), str(group)[:26], font=_font(17), fill="#93acbd")
                draw.rounded_rectangle((320, y, 1020, y + 25), 5, fill="#10283a")
                width = int(700 * abs(numeric) / maximum)
                color = "#ff667d" if numeric < 0 else palette[index]
                draw.rounded_rectangle((320, y, 320 + width, y + 25), 5, fill=color)
                unit = f" {metric_dataset['unit']}" if metric_dataset["unit"] else ""
                draw.text(
                    (1040, y - 2),
                    f"{value}{unit}",
                    font=_font(20, bold=True),
                    fill="#f2fbff",
                )
        draw.text(
            (540 if chart_type == "pie" else 82, 622),
            f"EVIDENCE: {metric_dataset['row_count']} ROWS / {str(metric_dataset['source_reference'])[:72]}",
            font=_font(13, mono=True),
            fill="#38d8ff",
        )
        footer = snapshot["document_footer"] or snapshot["action_boundary"]
        draw.text((82, 650), footer[:130], font=_font(14), fill="#7890a3")
        buffer = io.BytesIO()
        image.save(buffer, format="PNG", optimize=True)
        return buffer.getvalue()

    totals = snapshot["totals"]
    values = [totals["ALLOW"], totals["REVIEW"], totals["BLOCK"]]
    palette = ["#39e6a0", "#ffc857", "#ff667d"]
    labels = ["PASSED GATE", "NEEDS ATTENTION", "SAFELY STOPPED"]
    total = max(1, totals["TOTAL"])
    if chart_type == "pie":
        box = (100, 230, 500, 630)
        start = -90
        for value, color in zip(values, palette, strict=True):
            end = start + (360 * value / total)
            draw.pieslice(box, start=start, end=end, fill=color)
            start = end
        draw.ellipse((200, 330, 400, 530), fill="#07131f")
        total_text = str(totals["TOTAL"])
        bbox = draw.textbbox((0, 0), total_text, font=_font(58, bold=True))
        draw.text(
            (300 - (bbox[2] - bbox[0]) / 2, 370),
            total_text,
            font=_font(58, bold=True),
            fill="#f2fbff",
        )
        draw.text((256, 446), "DECISIONS", font=_font(14, mono=True), fill="#8ba2b4")
    else:
        for index, (value, color) in enumerate(zip(values, palette, strict=True)):
            y = 280 + index * 105
            draw.rounded_rectangle((105, y, 500, y + 38), 7, fill="#10283a")
            draw.rounded_rectangle(
                (105, y, 105 + int(395 * value / total), y + 38), 7, fill=color
            )
            draw.text(
                (105, y - 28), labels[index], font=_font(15, mono=True), fill="#8ba2b4"
            )
            draw.text(
                (516, y + 3), str(value), font=_font(25, bold=True), fill="#f2fbff"
            )

    for index, (label, value, color) in enumerate(
        zip(labels, values, palette, strict=True)
    ):
        y = 272 + index * 92
        draw.rounded_rectangle(
            (625, y, 1160, y + 67), 10, fill="#0c2232", outline="#173d54"
        )
        draw.ellipse((648, y + 21, 672, y + 45), fill=color)
        draw.text((692, y + 13), label, font=_font(16, mono=True), fill="#93acbd")
        draw.text((1090, y + 10), str(value), font=_font(30, bold=True), fill="#f2fbff")
    draw.text((625, 566), "IMPORTANT DETAIL", font=_font(14, mono=True), fill="#38d8ff")
    draw.text(
        (625, 596),
        snapshot["important_detail"],
        font=_font(23, bold=True),
        fill="#f2fbff",
    )
    footer = snapshot["document_footer"] or snapshot["action_boundary"]
    draw.text((82, 650), footer[:130], font=_font(14), fill="#7890a3")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def _set_cell_shading(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(column_width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            margins = properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                properties.append(margins)
            for edge, value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(108, 125, 139)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def _docx_report(snapshot: dict[str, Any], chart: bytes) -> bytes:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    document.core_properties.title = "ContextGate Evidence Report"
    document.core_properties.subject = "Evidence-grounded decision summary"
    document.core_properties.author = "ContextGate"

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = (
        f"{snapshot['company_name'].upper()}  /  CONTEXTGATE EVIDENCE REPORT"
        if snapshot["show_company_header"]
        else "CONTEXTGATE  /  EVIDENCE REPORT"
    )
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor(24, 119, 156)
    logo_content = _logo_bytes(snapshot)
    if logo_content:
        header.add_run("   ").add_picture(io.BytesIO(logo_content), width=Inches(0.42))
    footer = section.footer.paragraphs[0]
    if snapshot["document_footer"]:
        footer.text = snapshot["document_footer"]
        footer.runs[0].font.name = "Calibri"
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(108, 125, 139)
        footer = section.footer.add_paragraph()
    _add_page_field(footer)

    if snapshot["show_company_header"]:
        brand = document.add_paragraph()
        brand.paragraph_format.space_after = Pt(3)
        run = brand.add_run(snapshot["company_name"])
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(11, 37, 69)
        if snapshot["company_website"]:
            website = document.add_paragraph()
            website.paragraph_format.space_after = Pt(4)
            run = website.add_run(snapshot["company_website"])
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(24, 119, 156)
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("DECISION INTELLIGENCE BRIEF")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(24, 153, 198)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ContextGate Evidence Report")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        f"Operator: {snapshot['operator_name']}  |  {snapshot['generated_at'][:19]} UTC"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(82, 100, 113)

    callout = document.add_paragraph()
    callout.paragraph_format.space_after = Pt(12)
    run = callout.add_run(snapshot["action_boundary"])
    run.bold = True
    run.font.color.rgb = RGBColor(22, 101, 75)

    metrics = document.add_table(rows=2, cols=4)
    metrics.style = "Table Grid"
    headings = ("TOTAL", "PASSED GATE", "NEEDS ATTENTION", "SAFELY STOPPED")
    values = (
        snapshot["totals"]["TOTAL"],
        snapshot["totals"]["ALLOW"],
        snapshot["totals"]["REVIEW"],
        snapshot["totals"]["BLOCK"],
    )
    colors_by_column = ("0B2545", "146B4A", "8A6200", "8F2538")
    for index, heading in enumerate(headings):
        metrics.cell(0, index).text = heading
        metrics.cell(1, index).text = str(values[index])
        _set_cell_shading(metrics.cell(0, index), colors_by_column[index])
        _set_cell_shading(metrics.cell(1, index), "F2F6F8")
        for run in metrics.cell(0, index).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8.5)
        for run in metrics.cell(1, index).paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(colors_by_column[index])
            run.font.bold = True
            run.font.size = Pt(20)
        metrics.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        metrics.cell(1, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_geometry(metrics, [2340, 2340, 2340, 2340])

    document.add_heading("Decision signal", level=1)
    document.add_picture(io.BytesIO(chart), width=Inches(6.35))
    caption = document.add_paragraph("Figure 1. Resolved decision distribution.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(82, 100, 113)

    document.add_heading("Company context", level=1)
    for label, value in (
        ("Important detail", snapshot["important_detail"]),
        ("Identity fields", ", ".join(snapshot["identity_fields"]) or "Not configured"),
        ("Policy version", snapshot["policy_version"]),
    ):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(str(value))

    if snapshot["grouped_metrics"]:
        document.add_heading("Grouped metrics and evidence", level=1)
        for dataset in snapshot["grouped_metrics"]:
            document.add_heading(dataset["dataset_name"], level=2)
            descriptor = document.add_paragraph()
            descriptor.add_run(
                f"{dataset['metric_field']} by {dataset['group_field']} · "
            ).bold = True
            descriptor.add_run(
                f"{dataset['row_count']} contributing rows · "
                f"{dataset['source_reference']}"
            )
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, heading in enumerate(("Group", "Total", "Evidence rows")):
                table.cell(0, index).text = heading
                _set_cell_shading(table.cell(0, index), "E8EEF5")
                for run in table.cell(0, index).paragraphs[0].runs:
                    run.font.bold = True
            _set_repeat_header(table.rows[0])
            for group, total in dataset["group_totals"].items():
                references = [
                    str(item["reference"])
                    for item in dataset["evidence"]
                    if item["group"] == group
                ]
                cells = table.add_row().cells
                cells[0].text = group
                cells[
                    1
                ].text = f"{total}{' ' + dataset['unit'] if dataset['unit'] else ''}"
                cells[2].text = "\n".join(references) or "Not supplied"
            _set_table_geometry(table, [2100, 1500, 5760])

    document.add_heading("Cases and effective outcomes", level=1)
    cases = document.add_table(rows=1, cols=4)
    cases.style = "Table Grid"
    headers = ("Case", "Outcome", "Source", "Explanation")
    for index, heading in enumerate(headers):
        cases.cell(0, index).text = heading
        _set_cell_shading(cases.cell(0, index), "E8EEF5")
        for run in cases.cell(0, index).paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(11, 37, 69)
    _set_repeat_header(cases.rows[0])
    for item in snapshot["cases"]:
        cells = cases.add_row().cells
        cells[0].text = f"{item['case_id']}\n{item['title']}"
        cells[1].text = item["effective_outcome"]
        if item["corrected"]:
            cells[1].text += f"\nOriginal: {item['original_outcome']}"
        cells[2].text = item["source"]
        cells[3].text = item["summary"]
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    _set_table_geometry(cases, [1300, 1200, 2350, 4510])

    document.add_heading("Observed patterns", level=1)
    if snapshot["patterns"]:
        for item in snapshot["patterns"]:
            paragraph = document.add_paragraph()
            paragraph.add_run(
                f"{_text(item.get('label'), 'Pattern', 120)}: "
            ).bold = True
            paragraph.add_run(
                f"{_text(item.get('count'), '—', 60)} - "
                f"{_text(item.get('description'), 'No detail supplied.', 400)}"
            )
    else:
        document.add_paragraph("No patterns are available in the current view.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_report(snapshot: dict[str, Any], chart: bytes) -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.68 * inch,
        title="ContextGate Evidence Report",
        author="ContextGate",
    )
    base = getSampleStyleSheet()
    title = ParagraphStyle(
        "ContextGateTitle",
        parent=base["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=28,
        textColor=colors.HexColor("#0B2545"),
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    subtitle = ParagraphStyle(
        "ContextGateSubtitle",
        parent=base["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#526471"),
        spaceAfter=12,
    )
    heading = ParagraphStyle(
        "ContextGateHeading",
        parent=base["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#18779C"),
        spaceBefore=12,
        spaceAfter=7,
    )
    body = ParagraphStyle(
        "ContextGateBody",
        parent=base["BodyText"],
        fontName="Helvetica",
        fontSize=9.2,
        leading=13,
        textColor=colors.HexColor("#182A36"),
        spaceAfter=6,
    )
    small = ParagraphStyle(
        "ContextGateSmall",
        parent=body,
        fontSize=7.2,
        leading=9,
        spaceAfter=0,
    )

    def escaped(value: object) -> str:
        return html.escape(_text(value, "—", 2_000))

    story: list[Any] = []
    logo_content = _logo_bytes(snapshot)
    if logo_content:
        with Image.open(io.BytesIO(logo_content)) as logo_image:
            logo_width, logo_height = logo_image.size
        scale = min(0.72 / max(logo_width, 1), 0.54 / max(logo_height, 1))
        story.append(
            ReportLabImage(
                io.BytesIO(logo_content),
                width=logo_width * scale * inch,
                height=logo_height * scale * inch,
            )
        )
    if snapshot["show_company_header"]:
        story.append(
            Paragraph(
                escaped(snapshot["company_name"]),
                ParagraphStyle(
                    "CompanyBrand",
                    parent=body,
                    textColor=colors.HexColor("#0B2545"),
                    fontName="Helvetica-Bold",
                    fontSize=12,
                    leading=15,
                    spaceAfter=3,
                ),
            )
        )
        if snapshot["company_website"]:
            story.append(
                Paragraph(
                    escaped(snapshot["company_website"]),
                    ParagraphStyle(
                        "CompanyWebsite",
                        parent=body,
                        textColor=colors.HexColor("#18779C"),
                        fontSize=8.5,
                        leading=11,
                        spaceAfter=4,
                    ),
                )
            )
    story.extend(
        [
            Paragraph(
                "DECISION INTELLIGENCE BRIEF",
                ParagraphStyle(
                    "Kicker",
                    parent=body,
                    textColor=colors.HexColor("#1899C6"),
                    fontName="Helvetica-Bold",
                    fontSize=8.5,
                    leading=11,
                    spaceAfter=3,
                ),
            ),
            Paragraph("ContextGate Evidence Report", title),
            Paragraph(
                f"Operator: {escaped(snapshot['operator_name'])} &nbsp; | &nbsp; "
                f"{escaped(snapshot['generated_at'][:19])} UTC",
                subtitle,
            ),
            Paragraph(
                f"<b>{escaped(snapshot['action_boundary'])}</b>",
                ParagraphStyle(
                    "Boundary",
                    parent=body,
                    textColor=colors.HexColor("#16654B"),
                    backColor=colors.HexColor("#E9F7F1"),
                    borderPadding=8,
                    spaceAfter=13,
                ),
            ),
        ]
    )
    metrics_data = [
        ["TOTAL", "PASSED GATE", "NEEDS ATTENTION", "SAFELY STOPPED"],
        [
            str(snapshot["totals"]["TOTAL"]),
            str(snapshot["totals"]["ALLOW"]),
            str(snapshot["totals"]["REVIEW"]),
            str(snapshot["totals"]["BLOCK"]),
        ],
    ]
    metrics = Table(
        metrics_data, colWidths=[1.63 * inch] * 4, rowHeights=[0.28 * inch, 0.44 * inch]
    )
    metrics.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#0B2545")),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#146B4A")),
                ("BACKGROUND", (2, 0), (2, 0), colors.HexColor("#8A6200")),
                ("BACKGROUND", (3, 0), (3, 0), colors.HexColor("#8F2538")),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F2F6F8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("TEXTCOLOR", (0, 1), (0, 1), colors.HexColor("#0B2545")),
                ("TEXTCOLOR", (1, 1), (1, 1), colors.HexColor("#146B4A")),
                ("TEXTCOLOR", (2, 1), (2, 1), colors.HexColor("#8A6200")),
                ("TEXTCOLOR", (3, 1), (3, 1), colors.HexColor("#8F2538")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 7.2),
                ("FONTSIZE", (0, 1), (-1, 1), 19),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#B9C9D2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D7E2E8")),
            ]
        )
    )
    story.extend(
        [
            metrics,
            Paragraph("Decision signal", heading),
            ReportLabImage(io.BytesIO(chart), width=6.5 * inch, height=3.65625 * inch),
            Spacer(1, 5),
            Paragraph(
                f"<b>Important detail:</b> {escaped(snapshot['important_detail'])}<br/>"
                f"<b>Identity fields:</b> {escaped(', '.join(snapshot['identity_fields']) or 'Not configured')}<br/>"
                f"<b>Data origin:</b> {'Fictional demo' if snapshot['fictional'] else 'Scanned/uploaded sources'}",
                body,
            ),
        ]
    )
    if snapshot["grouped_metrics"]:
        story.append(Paragraph("Grouped metrics and evidence", heading))
        for dataset in snapshot["grouped_metrics"]:
            story.append(
                Paragraph(
                    f"<b>{escaped(dataset['dataset_name'])}</b><br/>"
                    f"{escaped(dataset['metric_field'])} by {escaped(dataset['group_field'])} · "
                    f"{dataset['row_count']} contributing rows · "
                    f"{escaped(dataset['source_reference'])}",
                    body,
                )
            )
            metric_rows: list[list[Any]] = [
                [
                    Paragraph("GROUP", small),
                    Paragraph("TOTAL", small),
                    Paragraph("EVIDENCE ROWS", small),
                ]
            ]
            for group, total in dataset["group_totals"].items():
                references = [
                    item["reference"]
                    for item in dataset["evidence"]
                    if item["group"] == group
                ]
                unit = f" {dataset['unit']}" if dataset["unit"] else ""
                metric_rows.append(
                    [
                        Paragraph(escaped(group), small),
                        Paragraph(escaped(f"{total}{unit}"), small),
                        Paragraph(
                            escaped("; ".join(references) or "Not supplied"), small
                        ),
                    ]
                )
            metric_table = Table(
                metric_rows,
                colWidths=[1.35 * inch, 0.9 * inch, 4.25 * inch],
                repeatRows=1,
                hAlign="LEFT",
            )
            metric_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B9C9D2")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(metric_table)
    story.append(Paragraph("Cases and effective outcomes", heading))
    case_rows: list[list[Any]] = [
        [
            Paragraph("CASE", small),
            Paragraph("OUTCOME", small),
            Paragraph("SOURCE", small),
            Paragraph("EXPLANATION", small),
        ]
    ]
    for item in snapshot["cases"]:
        outcome = escaped(item["effective_outcome"])
        if item["corrected"]:
            outcome += f"<br/><font size='6.5'>Original: {escaped(item['original_outcome'])}</font>"
        case_rows.append(
            [
                Paragraph(
                    f"<b>{escaped(item['case_id'])}</b><br/>{escaped(item['title'])}",
                    small,
                ),
                Paragraph(outcome, small),
                Paragraph(escaped(item["source"]), small),
                Paragraph(escaped(item["summary"]), small),
            ]
        )
    case_table = Table(
        case_rows,
        colWidths=[0.95 * inch, 0.72 * inch, 1.55 * inch, 3.28 * inch],
        repeatRows=1,
        hAlign="LEFT",
    )
    case_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B9C9D2")),
            ]
        )
    )
    story.extend([case_table, Paragraph("Observed patterns", heading)])
    if snapshot["patterns"]:
        for item in snapshot["patterns"]:
            story.append(
                KeepTogether(
                    Paragraph(
                        f"<b>{escaped(item.get('label'))}: {escaped(item.get('count'))}</b><br/>"
                        f"{escaped(item.get('description'))}",
                        body,
                    )
                )
            )
    else:
        story.append(Paragraph("No patterns are available in the current view.", body))

    def decorate_page(canvas: Any, document_template: Any) -> None:
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#18779C"))
        canvas.setFont("Helvetica-Bold", 7.5)
        footer_text = snapshot["document_footer"] or "CONTEXTGATE / EVIDENCE REPORT"
        canvas.drawString(0.72 * inch, 0.4 * inch, footer_text[:95])
        canvas.setFillColor(colors.HexColor("#6C7D8B"))
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(
            7.78 * inch, 0.4 * inch, f"Page {document_template.page}"
        )
        canvas.restoreState()

    document.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)
    return buffer.getvalue()


def _html_report(snapshot: dict[str, Any]) -> bytes:
    rows = "".join(
        "<tr>"
        f"<td>{html.escape(item['case_id'])}</td>"
        f"<td>{html.escape(item['title'])}</td>"
        f"<td>{html.escape(item['effective_outcome'])}</td>"
        f"<td>{html.escape(item['source'])}</td>"
        f"<td>{html.escape(item['summary'])}</td>"
        "</tr>"
        for item in snapshot["cases"]
    )
    metric_sections = "".join(
        "<section class='metric'>"
        f"<h3>{html.escape(dataset['dataset_name'])}</h3>"
        f"<p>{html.escape(dataset['metric_field'])} by "
        f"{html.escape(dataset['group_field'])}; {dataset['row_count']} contributing "
        f"rows; source <code>{html.escape(dataset['source_reference'])}</code></p>"
        "<table><thead><tr><th>Group</th><th>Total</th><th>Evidence rows</th></tr></thead><tbody>"
        + "".join(
            "<tr>"
            f"<td>{html.escape(group)}</td>"
            f"<td>{html.escape(str(total))}{' ' + html.escape(dataset['unit']) if dataset['unit'] else ''}</td>"
            "<td>"
            + "<br>".join(
                html.escape(str(item["reference"]))
                for item in dataset["evidence"]
                if item["group"] == group
            )
            + "</td></tr>"
            for group, total in dataset["group_totals"].items()
        )
        + "</tbody></table></section>"
        for dataset in snapshot["grouped_metrics"]
    )
    logo = (
        f'<img class="logo" src="{html.escape(snapshot["company_logo_data_url"], quote=True)}" alt="Company logo">'
        if snapshot["company_logo_data_url"]
        else ""
    )
    website = (
        f'<p><a href="{html.escape(snapshot["company_website"], quote=True)}">{html.escape(snapshot["company_website"])}</a></p>'
        if snapshot["company_website"]
        else ""
    )
    brand = (
        f"<h2>{html.escape(snapshot['company_name'])}</h2>{website}"
        if snapshot["show_company_header"]
        else ""
    )
    footer = (
        f"<footer><p>{html.escape(snapshot['document_footer'])}</p></footer>"
        if snapshot["document_footer"]
        else ""
    )
    document = f"""<!doctype html><html lang="en"><head><meta charset="utf-8"><title>ContextGate Evidence Report</title><style>
body{{max-width:980px;margin:34px auto;padding:0 24px;font:15px/1.5 Arial,sans-serif;color:#142631}}header{{border-top:7px solid #18779c;padding-top:18px}}h1{{color:#0b2545}}.logo{{display:block;max-width:120px;max-height:72px;object-fit:contain}}.boundary{{padding:12px;background:#e9f7f1;color:#16654b;font-weight:700}}table{{width:100%;border-collapse:collapse;font-size:12px}}th,td{{padding:8px;border:1px solid #b9c9d2;text-align:left;vertical-align:top}}th{{background:#e8eef5;color:#0b2545}}footer{{margin-top:24px;color:#526471;border-top:1px solid #b9c9d2}}@media print{{body{{margin:0;max-width:none}}}}</style></head><body>
<header>{logo}{brand}<p>DECISION INTELLIGENCE BRIEF</p><h1>ContextGate Evidence Report</h1><p>Operator: {html.escape(snapshot["operator_name"])} | {html.escape(snapshot["generated_at"])}</p></header>
<p class="boundary">{html.escape(snapshot["action_boundary"])}</p><h2>Decision totals</h2><p>{snapshot["totals"]["ALLOW"]} passed | {snapshot["totals"]["REVIEW"]} need attention | {snapshot["totals"]["BLOCK"]} stopped | {snapshot["totals"]["TOTAL"]} total</p>
{f"<h2>Grouped metrics and row evidence</h2>{metric_sections}" if metric_sections else ""}<h2>Cases</h2><table><thead><tr><th>Case</th><th>Title</th><th>Outcome</th><th>Source</th><th>Explanation</th></tr></thead><tbody>{rows}</tbody></table>{footer}</body></html>"""
    return document.encode("utf-8")


def _requested_outputs(request: str) -> tuple[list[str], Literal["pie", "bar"]]:
    normalized = re.sub(r"[^a-z0-9]+", " ", request.casefold()).strip()
    chart_type: Literal["pie", "bar"] = "pie" if "pie" in normalized else "bar"
    wants_chart = any(word in normalized for word in ("chart", "graph", "image", "png"))
    wants_report = any(
        word in normalized
        for word in ("report", "pdf", "word", "docx", "document", "html")
    )
    if "combo" in normalized or "combination" in normalized:
        wants_chart = True
        wants_report = True
    outputs: list[str] = []
    if wants_report:
        explicit_docx = any(word in normalized for word in ("word", "docx", "doc"))
        explicit_pdf = "pdf" in normalized
        explicit_html = "html" in normalized
        if explicit_docx:
            outputs.append("docx")
        if explicit_pdf:
            outputs.append("pdf")
        if explicit_html:
            outputs.append("html")
        if not any((explicit_docx, explicit_pdf, explicit_html)):
            outputs.extend(("pdf", "docx"))
    if wants_chart:
        outputs.append("png")
    if not outputs:
        raise ReportExportError(
            "Ask for a report, PDF, Word document, graph, pie chart, or image."
        )
    return list(dict.fromkeys(outputs)), chart_type


def is_export_instruction(message: str) -> bool:
    """Return true only for an explicit create/save/export instruction."""

    normalized = re.sub(r"[^a-z0-9]+", " ", message.casefold()).strip()
    if re.search(
        r"\b(?:do not|don t|dont|never|not to)\s+"
        r"(?:create|make|generate|build|prepare|save|export)\b",
        normalized,
    ):
        return False
    has_action = any(
        re.search(rf"\b{verb}\b", normalized)
        for verb in (
            "create",
            "make",
            "generate",
            "build",
            "prepare",
            "save",
            "export",
        )
    )
    has_artifact = any(
        re.search(rf"\b{noun}\b", normalized)
        for noun in (
            "report",
            "pdf",
            "word",
            "docx",
            "document",
            "graph",
            "chart",
            "image",
            "png",
            "combo",
            "combination",
        )
    )
    return has_action and has_artifact


def create_exports(
    state: dict[str, Any],
    request: str,
    *,
    output_root: str | os.PathLike[str] | None = None,
) -> list[ExportArtifact]:
    """Create requested local artifacts and return exact file metadata."""

    if not isinstance(state, dict) or not isinstance(request, str):
        raise ReportExportError("The export request is invalid.")
    outputs, chart_type = _requested_outputs(request)
    root = _output_root(output_root)
    snapshot = _normalized_snapshot(state)
    normalized_request = re.sub(r"[^a-z0-9]+", " ", request.casefold()).strip()
    if any(
        phrase in normalized_request
        for phrase in (
            "without company name",
            "without the company name",
            "no company name",
            "do not add the company name",
            "do not add company name",
            "do not show the company name",
            "do not include the company name",
        )
    ):
        snapshot["show_company_header"] = False
    timestamp = datetime.now(UTC).strftime("%Y%m%d-%H%M%S")
    company_slug = "workspace"
    if snapshot["show_company_header"]:
        company_slug = (
            re.sub(r"[^a-z0-9]+", "-", snapshot["company_name"].casefold()).strip("-")[
                :40
            ]
            or "company"
        )
    stem = f"contextgate-{company_slug}-{timestamp}-{uuid4().hex[:6]}"
    chart = _chart_png(snapshot, chart_type)
    artifacts: list[ExportArtifact] = []
    builders = {
        "png": (f"{stem}-{chart_type}-chart.png", chart),
        "pdf": (f"{stem}-report.pdf", lambda: _pdf_report(snapshot, chart)),
        "docx": (f"{stem}-report.docx", lambda: _docx_report(snapshot, chart)),
        "html": (f"{stem}-report.html", lambda: _html_report(snapshot)),
    }
    try:
        for kind in outputs:
            filename, builder = builders[kind]
            content = builder() if callable(builder) else builder
            target = root / filename
            _atomic_write(target, content)
            artifacts.append(
                ExportArtifact(
                    kind=kind,
                    path=str(target),
                    filename=filename,
                    size_bytes=target.stat().st_size,
                )
            )
    except (OSError, ValueError, TypeError, KeyError):
        raise ReportExportError(
            "ContextGate could not create the requested local report safely."
        ) from None
    return artifacts
